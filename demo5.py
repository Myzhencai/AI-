import streamlit as st
from PIL import Image
import google.generativeai as genai
import docx
import PyPDF2
import pandas as pd
import io
import os
from docx import Document
from docx.shared import Pt, Inches
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# ============ 函数：生成格式化 docx ============
# def generate_structured_docx(text):
#     doc = Document()
#     lines = text.split('\n')
#
#     for line in lines:
#         line = line.strip()
#         if not line:
#             doc.add_paragraph("")  # 空行
#             continue
#
#         # 一级标题（如 1. 前置条件）
#         if line[:1].isdigit() and '.' in line[:4]:
#             para = doc.add_paragraph()
#             run = para.add_run(line)
#             run.bold = True
#             run.font.size = Pt(14)
#             continue
#
#         # 二级标题（如 - 背景介绍）
#         if line.startswith("- "):
#             para = doc.add_paragraph()
#             para.paragraph_format.left_indent = Inches(0.3)
#             run = para.add_run(line)
#             run.bold = True
#             continue
#
#         # 正文内容
#         para = doc.add_paragraph(line)
#         para.paragraph_format.left_indent = Inches(0.5)
#
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer



def generate_structured_docx(text):
    doc = Document()

    # 设置默认字体（仿宋/Garamond等更适合正式文档）
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    lines = text.split('\n')

    for idx, line in enumerate(lines):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # 一级编号标题（如 "1. 前置条件"）
        if line[:1].isdigit() and line[1:3] == '. ':
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(14)
            continue

        # 二级编号标题（如 "1.1 背景介绍"）
        if line[:3].count('.') == 1 and line[0].isdigit() and line[2].isdigit():
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            continue

        # 项目符号列表（如 “- xxx” 或 “• xxx”）
        if line.startswith(('- ', '• ', '○ ')):
            para = doc.add_paragraph(line[2:], style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.5)
            continue

        # 术语解释：像 “Logo: 品牌标识”
        if "：" in line or ":" in line:
            parts = line.split("：") if "：" in line else line.split(":")
            if len(parts) == 2:
                term, desc = parts
                para = doc.add_paragraph()
                run = para.add_run(f"{term.strip()}：")
                run.bold = True
                para.add_run(f"{desc.strip()}")
                para.paragraph_format.left_indent = Inches(0.6)
                continue

        # 正常段落
        para = doc.add_paragraph(line)
        para.paragraph_format.left_indent = Inches(0.4)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============ Prompt ============
roleprompt = """
我是一个产品需求文档分析专家，请提供一个PRD文档草稿给我，我会按照下面的 PRD 规范对内容进行结构化填充。
规范包括：

1. 前置条件
   - 背景介绍
   - 产品类目
   - 名词解释

2. 功能需求
   - 功能清单
   - 数据指标
   - 流程图（如 UI 框图）
   - 应用场景（使用场景、场景规则、边界判断、中断处理、功能与 UI 交互）
   - 结构图

3. 非功能说明
   - 性能指标（速度、可靠性、CPU/内存占用等）
   - 兼容性
   - 安全和保密

4. 验收标准

请提供给我一个PRD文档吧
"""

# ============ 页面设置 ============
st.set_page_config(page_title="Gemini Pro PRD分析助手", page_icon="♊")
st.title("🔧 零缺陷 PRD 分析 Agent")

st.write("欢迎使用 PRD 分析助手，请先输入 Google API 密钥👇")

# ============ API 密钥 ============
with st.expander("🔐 输入 Google API 密钥"):
    google_api_key = st.text_input("Google API 密钥", key="google_api_key", type="password")

if not google_api_key:
    st.info("⚠️ 请输入 Google API 密钥以继续")
    st.stop()

genai.configure(api_key=google_api_key)

# ============ 模型选择与文件上传 ============
with st.sidebar:
    option = st.selectbox('选择模型', ('gemini-2.0-flash-lite',))

    if 'model' not in st.session_state or st.session_state.model != option:
        st.session_state.chat = genai.GenerativeModel(option).start_chat(history=[])
        st.session_state.model = option

    temperature = st.slider("温度 (temperature)", 0.0, 1.0, 0.5, step=0.05)
    max_token = st.number_input("最大输出 token 数", min_value=256, value=2048)

    gen_config = genai.types.GenerationConfig(
        max_output_tokens=max_token,
        temperature=temperature
    )

    st.divider()
    upload_file = st.file_uploader(
        "📁 上传 PRD 草稿文档 (.docx / .pdf / .xlsx)",
        type=["docx", "pdf", "xls", "xlsx"]
    )

    file_text = ""
    if upload_file:
        st.write(f"📄 上传文件：{upload_file.name}")
        if upload_file.type == "application/pdf":
            reader = PyPDF2.PdfReader(upload_file)
            file_text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

        elif upload_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(upload_file)
            file_text = "\n".join([para.text for para in doc.paragraphs])

        elif upload_file.type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
            df = pd.read_excel(upload_file)
            file_text = df.to_csv(index=False)

        # 只处理一次文件
        if st.session_state.get("uploaded_filename") != upload_file.name:
            st.session_state["file_processed"] = False
            st.session_state["uploaded_filename"] = upload_file.name

    if st.button("🧹 清除聊天历史"):
        st.session_state.messages = [{"role": "system", "content": roleprompt}]

# ============ 聊天区 ============
if "messages" not in st.session_state:
    st.session_state["messages"] = [{"role": "system", "content": roleprompt}]

for msg in st.session_state.messages:
    st.chat_message(msg["role"]).write(msg["content"])

if prompt := st.chat_input("请输入问题或请求分析文档："):
    st.chat_message("user").write(prompt)
    st.session_state.messages.append({"role": "user", "content": prompt})

    # 如果有上传文件且未处理，优先分析文件
    if file_text and not st.session_state.get("file_processed", False):
        full_prompt = f"{roleprompt}\n\n以下是用户上传的 PRD 文档内容：\n{file_text}"
        response = st.session_state.chat.send_message(full_prompt, stream=True, generation_config=gen_config)
        st.session_state["file_processed"] = True
    else:
        response = st.session_state.chat.send_message(prompt, stream=True, generation_config=gen_config)

    response.resolve()
    result_text = response.text
    st.chat_message("assistant").write(result_text)
    st.session_state.messages.append({"role": "assistant", "content": result_text})

# ============ 下载按钮 ============
last_assistant_msg = next((m["content"] for m in reversed(st.session_state.messages) if m["role"] == "assistant"), None)

if last_assistant_msg:
    docx_buffer = generate_structured_docx(last_assistant_msg)
    st.download_button(
        label="📄 下载结构化 PRD 文档（.docx）",
        data=docx_buffer,
        file_name="结构化PRD.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
