import streamlit as st
from PIL import Image
import google.generativeai as genai
import docx
import PyPDF2
import pandas as pd
import io
import os
from docx import Document

# ============ 函数：生成 .docx 缓冲区供下载 ============
def generate_docx_download_buffer(text):
    """将文本内容转为 docx 并返回 BytesIO 缓冲区用于下载"""
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ============ PRD 结构化提示词 ============
roleprompt = """
我是一个产品需求文档分析专家，请提供一个PRD文档草稿给我，我会按照下面的 PRD 规范对内容进行结构化填充。
规范包括：

1. 前置条件
   - 背景介绍
   - 产品类目
   - 名词解释

2. 功能需求
   - 功能清单
   - 数据指标
   - 流程图（如 UI 框图）
   - 应用场景（使用场景、场景规则、边界判断、中断处理、功能与 UI 交互）
   - 结构图

3. 非功能说明
   - 性能指标（速度、可靠性、CPU/内存占用等）
   - 兼容性
   - 安全和保密

4. 验收标准

请提供给我一个PRD文档吧
"""

# ============ 页面设置 ============
st.set_page_config(page_title="Gemini Pro with Streamlit", page_icon="♊")
st.title("零缺陷Agent")

st.write("欢迎来到 Gemini Pro 聊天机器人。您可以通过提供您的 Google API 密钥来继续。")

# ============ 侧边栏 API 密钥 ============
with st.expander("提供您的 Google API 密钥"):
    google_api_key = st.text_input("Google API 密钥", key="google_api_key", type="password")

if not google_api_key:
    st.info("请输入 Google API 密钥以继续")
    st.stop()

genai.configure(api_key=google_api_key)

# ============ 模型选择、参数设定 ============
with st.sidebar:
    option = st.selectbox('选择您的模型', ('gemini-2.0-flash-lite',))

    if 'model' not in st.session_state or st.session_state.model != option:
        st.session_state.chat = genai.GenerativeModel(option).start_chat(history=[])
        st.session_state.model = option

    st.write("在此处调整您的参数:")
    temperature = st.number_input("温度", min_value=0.0, max_value=1.0, value=0.5, step=0.01)
    max_token = st.number_input("最大输出令牌数", min_value=0, value=10000)
    gen_config = genai.types.GenerationConfig(max_output_tokens=max_token, temperature=temperature)

    st.divider()

    upload_file = st.file_uploader(
        "在此上传您的文档（支持 .docx, .pdf, .xls, .xlsx）",
        accept_multiple_files=False,
        type=["docx", "pdf", "xls", "xlsx"]
    )

    file_text = ""
    if upload_file:
        file_details = {
            "filename": upload_file.name,
            "filetype": upload_file.type,
            "filesize": upload_file.size
        }
        st.write("文件信息：", file_details)

        if upload_file.type == "application/pdf":
            reader = PyPDF2.PdfReader(upload_file)
            file_text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

        elif upload_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(upload_file)
            file_text = "\n".join([para.text for para in doc.paragraphs])

        elif upload_file.type in ["application/vnd.ms-excel",
                                  "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
            df = pd.read_excel(upload_file)
            file_text = df.to_csv(index=False)

        if st.session_state.get("uploaded_filename") != upload_file.name:
            st.session_state["file_processed"] = False
            st.session_state["uploaded_filename"] = upload_file.name

    st.divider()

    if st.button("清除聊天历史"):
        st.session_state.messages.clear()
        st.session_state["messages"] = [{"role": "system", "content": roleprompt}]

# ============ 聊天逻辑 ============
if "messages" not in st.session_state:
    st.session_state["messages"] = [{"role": "system", "content": roleprompt}]

for msg in st.session_state.messages:
    st.chat_message(msg["role"]).write(msg["content"])

if prompt := st.chat_input():
    st.session_state.messages.append({"role": "user", "content": prompt})
    st.chat_message("user").write(prompt)

    if file_text and not st.session_state.get("file_processed", False):
        full_prompt = f"{roleprompt}\n\n以下是用户上传的 PRD 文档内容：\n{file_text}"
        response = st.session_state.chat.send_message(full_prompt, stream=True, generation_config=gen_config)
        st.session_state["file_processed"] = True
    elif not file_text:
        msg = "⚠️ 请先上传文档，我才能根据 PRD 规范进行分析。"
        st.session_state.messages.append({"role": "assistant", "content": msg})
        st.chat_message("assistant").write(msg)
        st.stop()
    else:
        response = st.session_state.chat.send_message(prompt, stream=True, generation_config=gen_config)

    response.resolve()
    msg = response.text
    st.session_state.messages.append({"role": "assistant", "content": msg})
    st.chat_message("assistant").write(msg)

# ============ 下载 .docx ============
last_assistant_msg = None
for message in reversed(st.session_state.messages):
    if message["role"] == "assistant":
        last_assistant_msg = message["content"]
        break

if last_assistant_msg:
    docx_buffer = generate_docx_download_buffer(last_assistant_msg)

    st.download_button(
        label="📄 下载 Word 文档（.docx）",
        data=docx_buffer,
        file_name="修改后的PRD.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
